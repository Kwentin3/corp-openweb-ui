from __future__ import annotations

import base64
import hashlib
import io
import re
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from html.parser import HTMLParser
from typing import Protocol
from uuid import uuid4

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt

from stage2_stt.config import SttConfig
from stage2_stt.contracts import MessageDocxExportRequestV1, MessageDocxExportResultV1


FACTORY_REQUIRED = "DocxExportAdapterFactory.create is the only production DOCX renderer entrypoint"
FORBIDDEN = "Route handlers must not call python-docx directly"

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
MAX_FILENAME_STEM = 80
SCAN_PARTS = (
    "word/document.xml",
    "docProps/core.xml",
    "docProps/app.xml",
    "word/_rels/document.xml.rels",
    "_rels/.rels",
)
FORBIDDEN_MARKERS = (
    "authorization",
    "bearer ",
    "cookie",
    "set-cookie",
    "stage2_",
    "lemonfox",
    "openai_api_key",
    "x-stage2-internal-token",
    "raw_provider",
    "provider_payload",
    "verbose_json",
    "http://stage2-stt",
    "stage2-stt:8080",
    "transcript_with_speakers",
    "copy response",
    "regenerate response",
    "good response",
    "bad response",
    "continue response",
    "previous-message-sentinel",
    "next-message-sentinel",
    "toolbar-control-sentinel",
)
SEMANTIC_PROFILE = "semantic_chat_v1"
FORMAT_WARNING_DEGRADED = "message_docx_formatting_degraded"
ALLOWED_HTML_TAGS = {
    "a",
    "b",
    "blockquote",
    "br",
    "code",
    "div",
    "em",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "hr",
    "i",
    "li",
    "ol",
    "p",
    "pre",
    "span",
    "strong",
    "table",
    "tbody",
    "td",
    "tfoot",
    "th",
    "thead",
    "tr",
    "ul",
}
DANGEROUS_HTML_TAGS = {"button", "form", "iframe", "img", "input", "noscript", "script", "select", "style", "svg", "textarea"}
BLOCK_TAGS = {"blockquote", "div", "h1", "h2", "h3", "h4", "h5", "h6", "hr", "ol", "p", "pre", "table", "ul"}


class MessageDocxExportError(RuntimeError):
    def __init__(self, code: str, message: str) -> None:
        self.code = code
        self.message = message
        super().__init__(f"{code}: {message}")


@dataclass
class HtmlNode:
    tag: str
    attrs: dict[str, str] = field(default_factory=dict)
    children: list["HtmlNode | str"] = field(default_factory=list)


class _SemanticHtmlParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.root = HtmlNode("root")
        self._stack: list[HtmlNode] = [self.root]
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        normalized_tag = tag.lower()
        if normalized_tag in DANGEROUS_HTML_TAGS:
            self._skip_depth += 1
            return
        if self._skip_depth:
            return
        if normalized_tag not in ALLOWED_HTML_TAGS:
            return
        safe_attrs = _safe_html_attrs(normalized_tag, attrs)
        node = HtmlNode(normalized_tag, safe_attrs)
        self._stack[-1].children.append(node)
        if normalized_tag != "br":
            self._stack.append(node)

    def handle_startendtag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        self.handle_starttag(tag, attrs)
        self.handle_endtag(tag)

    def handle_endtag(self, tag: str) -> None:
        normalized_tag = tag.lower()
        if normalized_tag in DANGEROUS_HTML_TAGS and self._skip_depth:
            self._skip_depth -= 1
            return
        if self._skip_depth:
            return
        for index in range(len(self._stack) - 1, 0, -1):
            if self._stack[index].tag == normalized_tag:
                del self._stack[index:]
                return

    def handle_data(self, data: str) -> None:
        if self._skip_depth or not data:
            return
        self._stack[-1].children.append(data)


class DocxExportAdapter(Protocol):
    def render(self, request: MessageDocxExportRequestV1, *, exported_at: datetime) -> bytes:
        ...


class DocxExportAdapterFactory:
    def create(self) -> DocxExportAdapter:
        return PythonDocxExportAdapter()


class PythonDocxExportAdapter:
    def render(self, request: MessageDocxExportRequestV1, *, exported_at: datetime) -> bytes:
        document = Document()
        document.core_properties.title = "Exported Assistant Message"
        document.core_properties.subject = "OpenWebUI message-level DOCX export"
        document.core_properties.author = "STT v2 extension layer"
        document.add_heading("Exported Assistant Message", level=1)

        metadata = _metadata_rows(request, exported_at)
        if metadata:
            table = document.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            for label, value in metadata:
                cells = table.add_row().cells
                cells[0].text = label
                cells[1].text = value
            document.add_paragraph("")

        if request.options.formatting_profile == SEMANTIC_PROFILE:
            _render_semantic_message(document, request)
        else:
            source_text = request.message_markdown or request.message_text
            _render_simple_markdown(document, source_text)
        document.add_paragraph("")
        footer = document.add_paragraph()
        footer.add_run("Generated by STT v2 extension layer").italic = True

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()


class MessageDocxExportService:
    def __init__(
        self,
        *,
        config: SttConfig,
        adapter: DocxExportAdapter | None = None,
    ) -> None:
        self.config = config
        self.adapter = adapter or DocxExportAdapterFactory().create()

    def export(self, request: MessageDocxExportRequestV1) -> MessageDocxExportResultV1:
        self._validate_request(request)
        exported_at = datetime.now(timezone.utc)
        try:
            payload = self.adapter.render(request, exported_at=exported_at)
        except MessageDocxExportError:
            raise
        except Exception as exc:
            raise MessageDocxExportError(
                "message_docx_generation_failed",
                "DOCX generation failed",
            ) from exc

        max_docx_bytes = self.config.message_docx_max_docx_mb * 1024 * 1024
        if len(payload) > max_docx_bytes:
            raise MessageDocxExportError(
                "message_docx_message_too_large",
                "Generated DOCX exceeds the configured base64 delivery limit",
            )
        _scan_generated_docx(payload)

        checksum = hashlib.sha256(payload).hexdigest()
        return MessageDocxExportResultV1(
            export_id=f"docx-{uuid4().hex}",
            filename=_filename_for_request(request, exported_at),
            content_type=DOCX_CONTENT_TYPE,
            size_bytes=len(payload),
            checksum_sha256=checksum,
            delivery="base64",
            download_payload_base64=base64.b64encode(payload).decode("ascii"),
            download_url=None,
            file_id=None,
            warnings=_warnings_for_request(request),
        )

    def _validate_request(self, request: MessageDocxExportRequestV1) -> None:
        if request.message_role != "assistant":
            raise MessageDocxExportError(
                "message_docx_unsupported_role",
                "Only assistant messages can be exported to DOCX in MVP",
            )
        if not request.message_text or not request.message_text.strip():
            raise MessageDocxExportError(
                "message_docx_empty_message",
                "Selected assistant message is empty",
            )
        if len(request.message_text) > self.config.message_docx_max_message_chars:
            raise MessageDocxExportError(
                "message_docx_message_too_large",
                "Selected assistant message exceeds the configured DOCX export limit",
            )
        if not request.source:
            raise MessageDocxExportError(
                "message_docx_no_safe_source",
                "DOCX export request does not identify a safe message source",
            )
        source_path = request.safe_metadata.source_url_path
        if source_path and not source_path.startswith("/"):
            raise MessageDocxExportError(
                "message_docx_no_safe_source",
                "DOCX export source path must be a relative application path",
            )
        if _contains_forbidden_marker(request.message_text):
            raise MessageDocxExportError(
                "message_docx_no_leak_check_failed",
                "Selected assistant message contains forbidden internal markers",
            )
        if request.message_markdown and _contains_forbidden_marker(request.message_markdown):
            raise MessageDocxExportError(
                "message_docx_no_leak_check_failed",
                "Selected assistant message markdown contains forbidden internal markers",
            )
        if request.message_html and _contains_forbidden_marker(request.message_html):
            raise MessageDocxExportError(
                "message_docx_no_leak_check_failed",
                "Selected assistant message HTML contains forbidden internal markers",
            )


def _metadata_rows(
    request: MessageDocxExportRequestV1,
    exported_at: datetime,
) -> list[tuple[str, str]]:
    metadata = request.safe_metadata
    rows = [("Export time", exported_at.isoformat())]
    if request.options.include_chat_title and metadata.chat_title:
        rows.append(("Chat title", _safe_metadata_value(metadata.chat_title)))
    if request.options.include_model_name and metadata.model_name:
        rows.append(("Model", _safe_metadata_value(metadata.model_name)))
    if request.options.include_timestamp and metadata.message_timestamp:
        rows.append(("Message time", _safe_metadata_value(metadata.message_timestamp)))
    if metadata.result_ref:
        rows.append(("Result reference", _safe_metadata_value(metadata.result_ref)))
    if metadata.source_url_path:
        rows.append(("Source", _safe_metadata_value(metadata.source_url_path)))
    return [(label, value) for label, value in rows if value]


def _render_simple_markdown(document: Document, text: str) -> None:
    in_code = False
    code_lines: list[str] = []
    for raw_line in str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.rstrip()
        if line.strip().startswith("```"):
            if in_code:
                _add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)) + 1, 4)
            document.add_heading(_strip_inline_markdown(heading.group(2)), level=level)
            continue
        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet:
            document.add_paragraph(_strip_inline_markdown(bullet.group(1)), style="List Bullet")
            continue
        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered:
            document.add_paragraph(
                _strip_inline_markdown(numbered.group(1)),
                style="List Number",
            )
            continue
        document.add_paragraph(_strip_inline_markdown(stripped))
    if in_code and code_lines:
        _add_code_block(document, code_lines)


def _render_semantic_message(document: Document, request: MessageDocxExportRequestV1) -> None:
    if request.message_html and request.message_html.strip():
        root = _parse_semantic_html(request.message_html)
        _render_html_children(document, root.children)
        return
    if request.message_markdown and _looks_structured_markdown(request.message_markdown):
        _render_semantic_markdown(document, request.message_markdown)
        return
    _render_simple_markdown(document, request.message_text)


def _parse_semantic_html(html: str) -> HtmlNode:
    parser = _SemanticHtmlParser()
    try:
        parser.feed(html)
        parser.close()
    except MessageDocxExportError:
        raise
    except Exception as exc:
        raise MessageDocxExportError(
            "message_docx_unsafe_html",
            "Selected assistant message HTML could not be parsed safely",
        ) from exc
    return parser.root


def _render_html_children(
    document: Document,
    children: list[HtmlNode | str],
    *,
    list_stack: list[str] | None = None,
) -> None:
    active_list_stack = list_stack or []
    pending_text: list[HtmlNode | str] = []
    for child in children:
        if isinstance(child, str):
            if child.strip():
                pending_text.append(child)
            continue
        if child.tag not in BLOCK_TAGS:
            pending_text.append(child)
            continue
        if pending_text:
            paragraph = document.add_paragraph()
            _add_inline_html(paragraph, pending_text)
            pending_text = []
        _render_html_block(document, child, list_stack=active_list_stack)
    if pending_text:
        paragraph = document.add_paragraph()
        _add_inline_html(paragraph, pending_text)


def _render_html_block(
    document: Document,
    node: HtmlNode,
    *,
    list_stack: list[str],
) -> None:
    if node.tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = min(int(node.tag[1]), 4)
        paragraph = document.add_heading("", level=level)
        _add_inline_html(paragraph, node.children)
        return
    if node.tag in {"p", "div"}:
        if _has_block_child(node):
            _render_html_children(document, node.children, list_stack=list_stack)
        else:
            paragraph = document.add_paragraph()
            _add_inline_html(paragraph, node.children)
        return
    if node.tag == "blockquote":
        paragraph = _add_safe_paragraph(document, style="Quote")
        _add_inline_html(paragraph, node.children)
        return
    if node.tag == "pre":
        _add_code_block(document, _node_plain_text(node).splitlines())
        return
    if node.tag in {"ul", "ol"}:
        _render_html_list(document, node, list_stack=[*list_stack, node.tag])
        return
    if node.tag == "table":
        _render_html_table(document, node)
        document.add_paragraph("")
        return
    if node.tag == "hr":
        document.add_paragraph("---")


def _render_html_list(document: Document, node: HtmlNode, *, list_stack: list[str]) -> None:
    ordered = node.tag == "ol"
    level = max(len(list_stack) - 1, 0)
    for child in node.children:
        if not isinstance(child, HtmlNode) or child.tag != "li":
            continue
        inline_children: list[HtmlNode | str] = []
        nested_lists: list[HtmlNode] = []
        for item_child in child.children:
            if isinstance(item_child, HtmlNode) and item_child.tag in {"ul", "ol"}:
                nested_lists.append(item_child)
            else:
                inline_children.append(item_child)
        paragraph = _add_safe_paragraph(
            document,
            style="List Number" if ordered else "List Bullet",
        )
        if level:
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)
        _add_inline_html(paragraph, inline_children)
        for nested in nested_lists:
            _render_html_list(document, nested, list_stack=[*list_stack, nested.tag])


def _render_html_table(document: Document, node: HtmlNode) -> None:
    rows = _table_rows(node)
    if not rows:
        return
    max_columns = max(len(row) for row in rows)
    if max_columns <= 0:
        return
    table = document.add_table(rows=len(rows), cols=max_columns)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index, cell_node in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            lines = [line for line in _node_plain_text(cell_node).splitlines() if line.strip()]
            if not lines:
                continue
            first = cell.paragraphs[0]
            first.text = lines[0]
            if cell_node.tag == "th":
                for run in first.runs:
                    run.bold = True
            for line in lines[1:]:
                cell.add_paragraph(line)


def _table_rows(node: HtmlNode) -> list[list[HtmlNode]]:
    rows: list[list[HtmlNode]] = []
    if node.tag == "tr":
        cells = [child for child in node.children if isinstance(child, HtmlNode) and child.tag in {"td", "th"}]
        if cells:
            rows.append(cells)
        return rows
    for child in node.children:
        if isinstance(child, HtmlNode):
            rows.extend(_table_rows(child))
    return rows


def _add_inline_html(
    paragraph,
    children: list[HtmlNode | str],
    *,
    bold: bool = False,
    italic: bool = False,
    code: bool = False,
) -> None:
    for child in children:
        if isinstance(child, str):
            _add_text_run(paragraph, _normalize_inline_text(child), bold=bold, italic=italic, code=code)
            continue
        if child.tag == "br":
            paragraph.add_run().add_break()
            continue
        if child.tag in {"strong", "b"}:
            _add_inline_html(paragraph, child.children, bold=True, italic=italic, code=code)
            continue
        if child.tag in {"em", "i"}:
            _add_inline_html(paragraph, child.children, bold=bold, italic=True, code=code)
            continue
        if child.tag == "code":
            _add_inline_html(paragraph, child.children, bold=bold, italic=italic, code=True)
            continue
        if child.tag == "a":
            href = child.attrs.get("href")
            label = _node_plain_text(child).strip() or href or ""
            if href and label:
                _add_hyperlink(paragraph, label, href)
            elif label:
                _add_text_run(paragraph, label, bold=bold, italic=italic, code=code)
            continue
        if child.tag in BLOCK_TAGS:
            _add_text_run(paragraph, _node_plain_text(child), bold=bold, italic=italic, code=code)
            continue
        _add_inline_html(paragraph, child.children, bold=bold, italic=italic, code=code)


def _add_text_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, code: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bold or None
    run.italic = italic or None
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def _add_hyperlink(paragraph, label: str, href: str) -> None:
    relationship_id = paragraph.part.relate_to(
        href,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    properties.append(style)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_safe_paragraph(document: Document, *, style: str | None = None):
    try:
        return document.add_paragraph(style=style)
    except KeyError:
        return document.add_paragraph()


def _node_plain_text(node: HtmlNode | str) -> str:
    if isinstance(node, str):
        return _normalize_inline_text(node)
    parts: list[str] = []
    for child in node.children:
        if isinstance(child, HtmlNode) and child.tag == "br":
            parts.append("\n")
        else:
            parts.append(_node_plain_text(child))
    value = "".join(parts)
    if node.tag in {"p", "div", "li", "tr"}:
        value = value.strip()
    return value


def _has_block_child(node: HtmlNode) -> bool:
    return any(isinstance(child, HtmlNode) and child.tag in BLOCK_TAGS for child in node.children)


def _normalize_inline_text(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "").replace("\u00a0", " "))


def _safe_html_attrs(tag: str, attrs: list[tuple[str, str | None]]) -> dict[str, str]:
    safe: dict[str, str] = {}
    for raw_name, raw_value in attrs:
        name = raw_name.lower()
        value = raw_value or ""
        if name.startswith("on"):
            raise MessageDocxExportError(
                "message_docx_unsafe_html",
                "Selected assistant message HTML contains event handlers",
            )
        if tag == "a" and name == "href":
            href = _safe_href(value)
            if href:
                safe["href"] = href
    return safe


def _safe_href(value: str) -> str | None:
    href = str(value or "").strip()
    if not href:
        return None
    lowered = href.lower()
    if lowered.startswith(("javascript:", "data:", "file:", "vbscript:")):
        raise MessageDocxExportError(
            "message_docx_unsafe_html",
            "Selected assistant message HTML contains an unsafe link",
        )
    if lowered.startswith(("http://", "https://", "mailto:")):
        return href
    return None


def _looks_structured_markdown(text: str | None) -> bool:
    value = str(text or "")
    patterns = (
        r"^#{1,6}\s+",
        r"^\s*[-*+]\s+",
        r"^\s*\d+[.)]\s+",
        r"```",
        r"\|.+\|",
        r"\*\*[^*]+\*\*",
        r"\[[^\]]+\]\([^)]+\)",
        r"^>\s+",
    )
    return any(re.search(pattern, value, flags=re.MULTILINE) for pattern in patterns)


def _render_semantic_markdown(document: Document, text: str) -> None:
    lines = str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()
        if stripped.startswith("```"):
            if in_code:
                _add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw_line.rstrip())
            index += 1
            continue
        if _is_markdown_table_start(lines, index):
            rows, next_index = _consume_markdown_table(lines, index)
            _render_markdown_table(document, rows)
            index = next_index
            continue
        if not stripped:
            document.add_paragraph("")
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            paragraph = document.add_heading("", level=min(len(heading.group(1)), 4))
            _add_inline_markdown(paragraph, heading.group(2))
            index += 1
            continue
        bullet = re.match(r"^(\s*)[-*+]\s+(.+)$", raw_line)
        if bullet:
            paragraph = _add_safe_paragraph(document, style="List Bullet")
            level = len(bullet.group(1).replace("\t", "  ")) // 2
            if level:
                paragraph.paragraph_format.left_indent = Inches(0.25 * level)
            _add_inline_markdown(paragraph, bullet.group(2))
            index += 1
            continue
        numbered = re.match(r"^(\s*)\d+[.)]\s+(.+)$", raw_line)
        if numbered:
            paragraph = _add_safe_paragraph(document, style="List Number")
            level = len(numbered.group(1).replace("\t", "  ")) // 2
            if level:
                paragraph.paragraph_format.left_indent = Inches(0.25 * level)
            _add_inline_markdown(paragraph, numbered.group(2))
            index += 1
            continue
        quote = re.match(r"^>\s+(.+)$", stripped)
        if quote:
            paragraph = _add_safe_paragraph(document, style="Quote")
            _add_inline_markdown(paragraph, quote.group(1))
            index += 1
            continue
        paragraph = document.add_paragraph()
        _add_inline_markdown(paragraph, stripped)
        index += 1
    if in_code and code_lines:
        _add_code_block(document, code_lines)


def _is_markdown_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return "|" in lines[index] and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[index + 1]) is not None


def _consume_markdown_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows = [_split_markdown_table_row(lines[index])]
    index += 2
    while index < len(lines) and "|" in lines[index] and lines[index].strip():
        rows.append(_split_markdown_table_row(lines[index]))
        index += 1
    return rows, index


def _split_markdown_table_row(line: str) -> list[str]:
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def _render_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            text = row[column_index] if column_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            _add_inline_markdown(paragraph, text)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph("")


def _add_inline_markdown(paragraph, text: str) -> None:
    pattern = re.compile(r"(\[([^\]]+)\]\(([^)]+)\)|`([^`]+)`|\*\*([^*]+)\*\*|\*([^*]+)\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            _add_text_run(paragraph, text[position : match.start()])
        if match.group(2) is not None:
            href = _safe_href(match.group(3))
            if href:
                _add_hyperlink(paragraph, match.group(2), href)
            else:
                _add_text_run(paragraph, match.group(2))
        elif match.group(4) is not None:
            _add_text_run(paragraph, match.group(4), code=True)
        elif match.group(5) is not None:
            _add_text_run(paragraph, match.group(5), bold=True)
        elif match.group(6) is not None:
            _add_text_run(paragraph, match.group(6), italic=True)
        position = match.end()
    if position < len(text):
        _add_text_run(paragraph, text[position:])


def _warnings_for_request(request: MessageDocxExportRequestV1) -> list[str]:
    if request.options.formatting_profile != SEMANTIC_PROFILE:
        return []
    if request.message_html and request.message_html.strip():
        return []
    if request.message_markdown and _looks_structured_markdown(request.message_markdown):
        return []
    return [FORMAT_WARNING_DEGRADED]


def _add_code_block(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def _strip_inline_markdown(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"\*([^*]+)\*", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", value)
    return value.strip()


def _filename_for_request(
    request: MessageDocxExportRequestV1,
    exported_at: datetime,
) -> str:
    seed = request.safe_metadata.chat_title or request.message_text.splitlines()[0]
    stem = _sanitize_filename_stem(seed) or "message-export"
    timestamp = exported_at.strftime("%Y%m%d-%H%M%S")
    return f"{stem}-{timestamp}.docx"


def _sanitize_filename_stem(value: str) -> str:
    stem = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', " ", str(value or ""))
    stem = re.sub(r"\s+", " ", stem).strip(" .")
    if len(stem) > MAX_FILENAME_STEM:
        stem = stem[:MAX_FILENAME_STEM].rstrip(" .")
    return stem


def _safe_metadata_value(value: str) -> str:
    cleaned = str(value or "").strip()
    if _contains_forbidden_marker(cleaned):
        raise MessageDocxExportError(
            "message_docx_no_leak_check_failed",
            "DOCX metadata contains forbidden internal markers",
        )
    return cleaned


def _scan_generated_docx(payload: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            for part in SCAN_PARTS:
                if part not in archive.namelist():
                    continue
                text = archive.read(part).decode("utf-8", errors="ignore")
                if _contains_forbidden_marker(text):
                    raise MessageDocxExportError(
                        "message_docx_no_leak_check_failed",
                        f"Generated DOCX contains forbidden marker in {part}",
                    )
    except zipfile.BadZipFile as exc:
        raise MessageDocxExportError(
            "message_docx_generation_failed",
            "Generated DOCX is not a valid ZIP package",
        ) from exc


def _contains_forbidden_marker(text: str) -> bool:
    lowered = str(text or "").lower()
    return any(marker in lowered for marker in FORBIDDEN_MARKERS)
