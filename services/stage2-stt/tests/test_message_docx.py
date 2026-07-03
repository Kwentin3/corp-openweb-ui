import base64
import hashlib
from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient

from stage2_stt.app import create_app
from stage2_stt.config import load_stt_config
from stage2_stt.contracts import MessageDocxExportRequestV1
from stage2_stt.message_docx import (
    FACTORY_REQUIRED,
    FORBIDDEN,
    MessageDocxExportError,
    MessageDocxExportService,
)


INTERNAL_TOKEN = "unit-test-internal-token"


def test_message_docx_export_generates_openable_docx_with_selected_message_only():
    selected_message = (
        "# Interview Notes\n\n"
        "Selected assistant message.\n\n"
        "- first point\n"
        "- second point\n\n"
        "```text\ncode block\n```\n"
    )
    request = MessageDocxExportRequestV1.model_validate(
        _request(
            message_text=selected_message,
            message_markdown=selected_message,
        )
    )
    result = MessageDocxExportService(config=load_stt_config({})).export(request)
    payload = base64.b64decode(result.download_payload_base64)
    document = Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert result.schema_version == "MessageDocxExportResultV1"
    assert result.content_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert result.delivery == "base64"
    assert result.download_url is None
    assert result.file_id is None
    assert result.checksum_sha256 == hashlib.sha256(payload).hexdigest()
    assert "Interview Notes" in text
    assert "Selected assistant message." in text
    assert "first point" in text
    assert "code block" in text
    assert "previous-message-sentinel" not in text
    assert "next-message-sentinel" not in text
    assert "toolbar-control-sentinel" not in text


def test_message_docx_semantic_html_preserves_visible_chat_structure():
    html = (
        "<h2>Interview Notes</h2>"
        "<p>Selected <strong>assistant</strong> <em>message</em> with "
        '<a href="https://example.com/source">source link</a>.</p>'
        "<ul><li>first point<ul><li>nested point</li></ul></li><li>second point</li></ul>"
        "<blockquote>quoted decision</blockquote>"
        "<table><thead><tr><th>Speaker</th><th>Decision</th></tr></thead>"
        "<tbody><tr><td>Speaker 1</td><td>Approve</td></tr></tbody></table>"
        "<pre><code>code block</code></pre>"
    )
    request = MessageDocxExportRequestV1.model_validate(
        _request(
            message_text="Interview Notes\nSelected assistant message.",
            message_markdown=None,
            message_html=html,
            options={
                "include_chat_title": True,
                "include_model_name": True,
                "include_timestamp": True,
                "formatting_profile": "semantic_chat_v1",
            },
        )
    )

    result = MessageDocxExportService(config=load_stt_config({})).export(request)
    payload = base64.b64decode(result.download_payload_base64)
    document = Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_rows = [
        [cell.text for cell in row.cells]
        for table in document.tables
        for row in table.rows
    ]
    hyperlink_targets = {
        rel.target_ref
        for rel in document.part.rels.values()
        if "hyperlink" in rel.reltype
    }

    assert result.warnings == []
    assert "Interview Notes" in text
    assert "first point" in text
    assert "nested point" in text
    assert "quoted decision" in text
    assert "code block" in text
    assert any(row == ["Speaker", "Decision"] for row in table_rows)
    assert any(row == ["Speaker 1", "Approve"] for row in table_rows)
    assert "https://example.com/source" in hyperlink_targets


def test_message_docx_semantic_markdown_precedes_truncated_html_source():
    markdown = (
        "## Meeting Notes\n\n"
        "Visible intro.\n\n"
        "---\n\n"
        "## Action items\n\n"
        "| Task | Owner | Due |\n"
        "|---|---|---|\n"
        "| Prepare letter | Olga | Soon |\n\n"
        "If needed, I can also format this for chat."
    )
    html = "<h2>Meeting Notes</h2><p>Visible intro.</p><hr><p>DOM-only truncated marker</p>"
    request = MessageDocxExportRequestV1.model_validate(
        _request(
            message_text="Meeting Notes\nVisible intro.",
            message_markdown=markdown,
            message_html=html,
            source="openwebui_chat_api",
            options={
                "include_chat_title": True,
                "include_model_name": True,
                "include_timestamp": True,
                "formatting_profile": "semantic_chat_v1",
            },
        )
    )

    result = MessageDocxExportService(config=load_stt_config({})).export(request)
    payload = base64.b64decode(result.download_payload_base64)
    document = Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_rows = [
        [cell.text for cell in row.cells]
        for table in document.tables
        for row in table.rows
    ]

    assert result.warnings == []
    assert "Action items" in text
    assert "If needed, I can also format this for chat." in text
    assert "DOM-only truncated marker" not in text
    assert any(row == ["Task", "Owner", "Due"] for row in table_rows)
    assert any(row == ["Prepare letter", "Olga", "Soon"] for row in table_rows)


def test_message_docx_semantic_profile_warns_when_structured_source_is_unavailable():
    request = MessageDocxExportRequestV1.model_validate(
        _request(
            message_text="Plain selected assistant message.",
            message_markdown=None,
            message_html=None,
            options={
                "include_chat_title": True,
                "include_model_name": True,
                "include_timestamp": True,
                "formatting_profile": "semantic_chat_v1",
            },
        )
    )

    result = MessageDocxExportService(config=load_stt_config({})).export(request)

    assert result.warnings == ["message_docx_formatting_degraded"]


def test_message_docx_semantic_profile_refuses_unsafe_html_links():
    request = MessageDocxExportRequestV1.model_validate(
        _request(
            message_text="Unsafe link",
            message_markdown=None,
            message_html='<p><a href="javascript:alert(1)">bad</a></p>',
            options={
                "include_chat_title": True,
                "include_model_name": True,
                "include_timestamp": True,
                "formatting_profile": "semantic_chat_v1",
            },
        )
    )
    service = MessageDocxExportService(config=load_stt_config({}))

    with pytest.raises(MessageDocxExportError) as exc_info:
        service.export(request)

    assert exc_info.value.code == "message_docx_unsafe_html"


def test_message_docx_export_refuses_non_assistant_role_with_typed_error():
    request = MessageDocxExportRequestV1.model_validate(_request(message_role="user"))
    service = MessageDocxExportService(config=load_stt_config({}))

    with pytest.raises(MessageDocxExportError) as exc_info:
        service.export(request)

    assert exc_info.value.code == "message_docx_unsupported_role"


def test_message_docx_export_refuses_empty_message_with_typed_error():
    request = MessageDocxExportRequestV1.model_validate(_request(message_text="  \n  "))
    service = MessageDocxExportService(config=load_stt_config({}))

    with pytest.raises(MessageDocxExportError) as exc_info:
        service.export(request)

    assert exc_info.value.code == "message_docx_empty_message"


def test_message_docx_export_refuses_too_large_message_with_typed_error():
    request = MessageDocxExportRequestV1.model_validate(_request(message_text="abcdef"))
    service = MessageDocxExportService(
        config=load_stt_config({"STAGE2_STT_MESSAGE_DOCX_MAX_MESSAGE_CHARS": "5"})
    )

    with pytest.raises(MessageDocxExportError) as exc_info:
        service.export(request)

    assert exc_info.value.code == "message_docx_message_too_large"


def test_message_docx_export_refuses_forbidden_markers_before_docx_return():
    request = MessageDocxExportRequestV1.model_validate(
        _request(message_text="Selected text\nprevious-message-sentinel")
    )
    service = MessageDocxExportService(config=load_stt_config({}))

    with pytest.raises(MessageDocxExportError) as exc_info:
        service.export(request)

    assert exc_info.value.code == "message_docx_no_leak_check_failed"


def test_message_docx_route_requires_internal_auth(monkeypatch):
    monkeypatch.setenv("STAGE2_STT_INTERNAL_API_KEY", INTERNAL_TOKEN)
    client = TestClient(create_app())

    response = client.post("/stage2-api/message-docx/exports", json=_request())

    assert response.status_code == 401
    assert response.json()["detail"]["code"] == "stage2_stt_internal_auth_failed"


def test_message_docx_route_returns_result_shape(monkeypatch):
    monkeypatch.setenv("STAGE2_STT_INTERNAL_API_KEY", INTERNAL_TOKEN)
    client = TestClient(create_app())

    response = client.post(
        "/stage2-api/message-docx/exports",
        headers={"Authorization": f"Bearer {INTERNAL_TOKEN}"},
        json=_request(),
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["schema_version"] == "MessageDocxExportResultV1"
    assert payload["delivery"] == "base64"
    assert payload["download_payload_base64"]
    assert payload["download_url"] is None
    assert payload["file_id"] is None


def test_message_docx_route_returns_typed_terminal_refusal(monkeypatch):
    monkeypatch.setenv("STAGE2_STT_INTERNAL_API_KEY", INTERNAL_TOKEN)
    monkeypatch.setenv("STAGE2_STT_MESSAGE_DOCX_MAX_MESSAGE_CHARS", "5")
    client = TestClient(create_app())

    response = client.post(
        "/stage2-api/message-docx/exports",
        headers={"Authorization": f"Bearer {INTERNAL_TOKEN}"},
        json=_request(message_text="too long"),
    )

    assert response.status_code == 413
    assert response.json()["detail"]["code"] == "message_docx_message_too_large"


def test_message_docx_route_returns_typed_unsafe_html_refusal(monkeypatch):
    monkeypatch.setenv("STAGE2_STT_INTERNAL_API_KEY", INTERNAL_TOKEN)
    client = TestClient(create_app())

    response = client.post(
        "/stage2-api/message-docx/exports",
        headers={"Authorization": f"Bearer {INTERNAL_TOKEN}"},
        json=_request(
            message_text="Unsafe link",
            message_markdown=None,
            message_html='<p><a href="javascript:alert(1)">bad</a></p>',
            options={
                "include_chat_title": True,
                "include_model_name": True,
                "include_timestamp": True,
                "formatting_profile": "semantic_chat_v1",
            },
        ),
    )

    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "message_docx_unsafe_html"


def test_message_docx_uses_factory_anchor():
    assert "DocxExportAdapterFactory.create" in FACTORY_REQUIRED
    assert "must not call python-docx directly" in FORBIDDEN


def _request(**overrides):
    payload = {
        "schema_version": "MessageDocxExportRequestV1",
        "request_id": "req-1",
        "chat_id": "chat-1",
        "message_id": "msg-1",
        "message_role": "assistant",
        "message_text": "Selected assistant message.",
        "message_markdown": None,
        "message_html": None,
        "source": "dom",
        "safe_metadata": {
            "chat_title": "Unit Chat",
            "model_name": "unit-model",
            "message_timestamp": "2026-07-03T00:00:00+00:00",
            "source_url_path": "/c/chat-1",
            "result_ref": "art_unit_reference",
        },
        "options": {
            "include_chat_title": True,
            "include_model_name": True,
            "include_timestamp": True,
            "formatting_profile": "simple_mvp",
        },
    }
    payload.update(overrides)
    return payload
